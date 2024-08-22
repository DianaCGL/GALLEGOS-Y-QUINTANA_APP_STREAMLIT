import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np
import io

# Definir las descripciones de las rúbricas específicas para cada pregunta
rubricas = {
    'Gestión de Acceso': {
        '¿Existen políticas y procedimientos documentados para la gestión de accesos?': {
            1: 'No se tienen políticas ni procedimientos documentados para la gestión de accesos. Esto implica que no hay controles formales para regular quién tiene acceso a qué información y sistemas, lo que deja la organización expuesta a riesgos significativos. Sin una documentación clara y procedimientos establecidos, no es posible auditar ni controlar el acceso de manera efectiva, lo que puede resultar en accesos no autorizados, pérdida o manipulación de datos críticos. La falta de estas políticas y procedimientos también indica una ausencia de conciencia y responsabilidad en la gestión de la seguridad de la información.',
            2: 'Existen políticas y procedimientos para la gestión de accesos, pero no están completamente documentados o actualizados. Esto puede llevar a inconsistencias en su aplicación y a la falta de claridad sobre quién tiene acceso a qué recursos y bajo qué circunstancias. Los documentos pueden estar desactualizados y no reflejan los cambios recientes en la estructura del sistema o en las responsabilidades del personal. Esto puede generar brechas de seguridad y procedimientos ineficaces, ya que las políticas no se alinean completamente con las prácticas actuales y las expectativas de la norma ISO 27001. Además, la falta de actualización regular sugiere que no se realizan revisiones periódicas necesarias para adaptarse a nuevos riesgos o cambios organizacionales.',
            3: 'Políticas y procedimientos para la gestión de accesos están documentados y se revisan regularmente. La mayoría de los requisitos de la norma ISO 27001 están cubiertos, incluyendo la asignación de roles y responsabilidades, la autorización de accesos y la revocación de privilegios cuando ya no son necesarios. Sin embargo, puede haber áreas donde las políticas no sean completamente específicas o no se hayan adaptado a todos los cambios tecnológicos recientes. A pesar de esto, los controles son adecuados y generalmente efectivos, aunque hay margen para mejorar la precisión y consistencia en algunos casos.',
            4: 'Las políticas y procedimientos de gestión de accesos cumplen completamente con los requisitos establecidos por la norma ISO 27001. Están totalmente documentados, actualizados y se revisan periódicamente para garantizar su eficacia. Incluyen procesos detallados para la gestión de cuentas de usuario, asignación de privilegios, monitoreo de accesos y respuesta a incidentes relacionados con accesos no autorizados. La documentación es clara y específica, y hay evidencia de su aplicación consistente. Además, se realiza una gestión proactiva para identificar y corregir posibles vulnerabilidades.',
            5: 'La implementación de políticas y procedimientos de gestión de accesos no solo cumple con todos los requisitos de la norma ISO 27001, sino que también incluye controles adicionales y medidas proactivas. Estos pueden incluir análisis de acceso basados en el comportamiento, herramientas avanzadas de monitoreo continuo y revisiones de accesos realizadas más frecuentemente que lo exigido por la norma. Se promueve una cultura de seguridad dentro de la organización, y el departamento de sistemas está constantemente buscando formas de mejorar la gestión de accesos a través de la capacitación continua del personal y la adopción de mejores prácticas y tecnologías emergentes.'
        },
        '¿Se implementan controles de autenticación fuertes para acceder a sistemas críticos?': {
            1: 'No se implementan controles de autenticación para acceder a sistemas críticos, lo que significa que cualquier persona podría potencialmente acceder a información y recursos sensibles sin ninguna verificación. Esto representa un grave riesgo para la seguridad de la información, ya que los sistemas críticos están expuestos a accesos no autorizados, lo cual podría resultar en pérdida de datos, alteración de información y otros incidentes de seguridad graves. La ausencia de controles de autenticación indica una falta de priorización de la seguridad en el acceso a los sistemas más importantes de la organización.',
            2: 'Se implementan controles de autenticación de manera limitada o inconsistente. Algunos sistemas críticos pueden tener autenticación básica, como contraseñas simples, pero faltan métodos robustos como la autenticación multifactor. La falta de una implementación consistente significa que hay variaciones en el nivel de seguridad entre diferentes sistemas, dejando brechas que pueden ser explotadas. Esto no cumple con los requisitos de ISO 27001, que exige controles de acceso adecuados y proporcionados a la sensibilidad de los sistemas y la información. Además, los procedimientos de autenticación existentes pueden no ser suficientes para resistir técnicas de ataque más avanzadas, como el phishing o el hacking.',
            3: 'Controles de autenticación fuertes están implementados de manera regular, cubriendo la mayoría de los sistemas críticos. Estos controles incluyen el uso de contraseñas complejas, autenticación multifactor y verificaciones periódicas de los accesos. Aunque los controles son adecuados y generalmente efectivos, pueden existir algunas áreas donde la implementación podría ser mejorada, como en la frecuencia de actualización de las contraseñas, la formación del personal en buenas prácticas de autenticación y la supervisión de los métodos de autenticación utilizados.',
            4: 'Los controles de autenticación cumplen totalmente con los requisitos de autenticación de ISO 27001. Todos los sistemas críticos están protegidos por métodos de autenticación robustos, como la autenticación multifactor, y hay políticas claras que regulan la gestión de contraseñas, la revocación de accesos y la supervisión de accesos. La autenticación es gestionada de manera centralizada y se realizan auditorías periódicas para garantizar la conformidad con las políticas. Los mecanismos de autenticación están diseñados para minimizar la posibilidad de errores humanos y están respaldados por procedimientos claros y bien documentados.',
            5: 'La implementación de controles de autenticación es avanzada y supera los requisitos estándar. Además de cumplir con ISO 27001, se utilizan tecnologías avanzadas como autenticación biométrica, tokens de hardware y análisis de comportamiento. Hay un monitoreo continuo de los accesos y un proceso de mejora continua para adaptar y fortalecer los métodos de autenticación basados en las últimas amenazas y mejores prácticas. Los usuarios reciben formación regular sobre la importancia de la seguridad en la autenticación y se fomenta una cultura de concienciación sobre la seguridad en toda la organización.'
        }
    },
    'Seguridad Física y Ambiental': {
        '¿Existen medidas de seguridad física para proteger los equipos críticos del departamento de sistemas?': {
            1: 'No hay medidas de seguridad física implementadas, dejando los equipos críticos vulnerables a accesos no autorizados y daños físicos. Esto incluye la falta de controles como cerraduras, sistemas de vigilancia y alarmas. La ausencia de medidas de seguridad física pone en riesgo la disponibilidad y la integridad de los sistemas y la información almacenada en ellos. Los equipos críticos pueden ser manipulados o dañados fácilmente, lo que puede resultar en interrupciones graves en las operaciones.',
            2: 'Las medidas de seguridad física existen pero son parciales o insuficientes. Puede haber algunas cerraduras o cámaras, pero no cubren todos los puntos de acceso o no están operativas en todo momento. Esto puede resultar en vulnerabilidades que pueden ser explotadas por intrusos o empleados no autorizados. Las medidas no cumplen completamente con los requisitos de ISO 27001, que exige una protección adecuada de los activos físicos. Además, la falta de controles físicos adecuados puede facilitar la introducción de dispositivos maliciosos en la red.',
            3: 'Las medidas de seguridad física están bien implementadas y protegen la mayoría de los equipos críticos. Esto incluye controles como acceso restringido a áreas sensibles, cámaras de vigilancia y alarmas operativas. Sin embargo, puede haber algunas áreas donde la seguridad física podría ser mejorada, como en la actualización de sistemas de seguridad obsoletos o en la implementación de controles adicionales en áreas críticas. También puede ser necesario realizar auditorías de seguridad física más frecuentes para asegurar que los controles están funcionando correctamente.',
            4: 'Las medidas de seguridad física cumplen totalmente con los requisitos de ISO 27001. Todos los equipos críticos están protegidos por controles físicos robustos, incluyendo acceso controlado mediante tarjetas de identificación, vigilancia constante y sistemas de alarma eficaces. Se realizan inspecciones periódicas para garantizar que los controles de seguridad física están operativos y efectivos. La documentación y los procedimientos están actualizados y reflejan las mejores prácticas en seguridad física.',
            5: 'La implementación de medidas de seguridad física es avanzada y supera los requisitos estándar. Además de los controles físicos básicos, se utilizan tecnologías avanzadas como sensores de movimiento, sistemas de autenticación biométrica y monitoreo continuo. La seguridad física es revisada y mejorada continuamente para adaptarse a nuevas amenazas y tecnologías emergentes. La organización también fomenta una cultura de concienciación sobre la seguridad física entre los empleados, proporcionando formación y recursos para mantener un entorno seguro.'
        },
        '¿Se realizan controles ambientales para proteger la infraestructura tecnológica (temperatura, humedad, etc.)?': {
            1: 'No se realizan controles ambientales, lo que puede llevar a fallos en la infraestructura tecnológica debido a condiciones ambientales adversas como altas temperaturas, humedad excesiva o fluctuaciones de energía. La falta de controles ambientales pone en riesgo la integridad y disponibilidad de los sistemas. Los equipos pueden sobrecalentarse, corroerse o sufrir daños eléctricos, lo que puede resultar en interrupciones del servicio y pérdida de datos.',
            2: 'Los controles ambientales existen pero se aplican de manera irregular o insuficiente. Puede haber sistemas de aire acondicionado, pero no están monitoreados o mantenidos adecuadamente. Esto puede resultar en condiciones ambientales que no son óptimas para la operación segura de la infraestructura tecnológica. La falta de mantenimiento regular y monitoreo continuo puede permitir que las condiciones ambientales se deterioren sin que nadie se dé cuenta, causando daños progresivos a los equipos.',
            3: 'Controles ambientales están bien implementados y protegen la mayoría de la infraestructura tecnológica. Esto incluye sistemas de control de temperatura, humedad y suministro de energía. Los controles se monitorean regularmente y se mantienen adecuadamente, aunque puede haber áreas donde se podrían implementar mejoras adicionales, como en la redundancia de sistemas críticos o en la integración de alarmas automáticas para alertar sobre condiciones adversas.',
            4: 'Los controles ambientales cumplen totalmente con los requisitos de ISO 27001, asegurando un ambiente controlado y seguro para toda la infraestructura tecnológica. Esto incluye sistemas de climatización y control de humedad, así como suministro de energía ininterrumpida. Se realizan inspecciones y mantenimientos periódicos para asegurar la eficacia de estos controles. La organización también tiene procedimientos documentados para la respuesta a incidentes ambientales y la recuperación de sistemas afectados.',
            5: 'La implementación de controles ambientales es avanzada y supera los requisitos estándar. Incluye sistemas de monitoreo continuo, alertas automatizadas y medidas preventivas adicionales como redundancia en sistemas críticos y técnicas avanzadas de control ambiental. La infraestructura tecnológica está protegida de manera óptima contra cualquier tipo de riesgo ambiental. Además, la organización está comprometida con la sostenibilidad y eficiencia energética, implementando prácticas verdes y tecnologías innovadoras para minimizar el impacto ambiental mientras se protege la infraestructura crítica.'
        }
    },
    'Gestión de Comunicaciones y Operaciones': {
        '¿Se utilizan procedimientos seguros para la transmisión de datos sensibles dentro y fuera de la organización?': {
            1: 'No se utilizan procedimientos seguros para la transmisión de datos, lo que expone la información sensible a interceptaciones y accesos no autorizados durante su transmisión. Esto incluye la falta de cifrado y de mecanismos de autenticación para proteger los datos en tránsito. La transmisión de datos sin protección puede resultar en la divulgación no autorizada de información confidencial, comprometiendo la seguridad de la organización y violando regulaciones de protección de datos.',
            2: 'Los procedimientos seguros para la transmisión de datos existen pero se utilizan de manera limitada o inconsistente. Algunos datos pueden ser cifrados, pero no todos, o el cifrado puede no ser robusto. Esto puede resultar en vulnerabilidades durante la transmisión de información sensible. Además, la falta de un enfoque consistente puede llevar a que los empleados utilicen métodos inseguros sin darse cuenta, lo que aumenta el riesgo de interceptación y compromiso de datos.',
            3: 'Procedimientos seguros para la transmisión de datos están bien implementados y se utilizan regularmente. Esto incluye el cifrado de datos en tránsito y la autenticación de las partes que participan en la comunicación. Sin embargo, puede haber áreas donde los procedimientos podrían ser mejorados, como en la implementación de métodos de cifrado más avanzados, la capacitación del personal sobre la importancia de estos procedimientos y la verificación periódica de que los métodos utilizados siguen siendo efectivos contra las amenazas actuales.',
            4: 'Los procedimientos seguros para la transmisión de datos cumplen totalmente con los requisitos de ISO 27001. Todos los datos sensibles que se transmiten dentro y fuera de la organización están protegidos por métodos robustos de cifrado y autenticación. Los procedimientos están documentados y se revisan regularmente para asegurar su efectividad. Además, la organización realiza auditorías periódicas para verificar la conformidad con estos procedimientos y para identificar posibles áreas de mejora.',
            5: 'La implementación de procedimientos seguros para la transmisión de datos es avanzada y supera los requisitos estándar. Esto incluye el uso de tecnologías de cifrado de última generación, monitoreo continuo de las transmisiones y la implementación de medidas adicionales como el cifrado de extremo a extremo y la verificación de integridad de los datos. La organización está proactivamente mejorando estos procedimientos para adaptarse a nuevas amenazas y tecnologías emergentes. También se fomenta una cultura de seguridad en la transmisión de datos entre los empleados, proporcionando formación continua y recursos para asegurar la adherencia a las mejores prácticas.'
        },
        '¿Se realizan pruebas periódicas de vulnerabilidades y evaluaciones de riesgos en la infraestructura de redes?': {
            1: 'No se realizan pruebas de vulnerabilidades ni evaluaciones de riesgos, lo que deja la infraestructura de redes expuesta a posibles amenazas y ataques. La falta de pruebas y evaluaciones significa que las vulnerabilidades pueden pasar desapercibidas y no se toman medidas para mitigarlas. Esto puede resultar en la explotación de fallos de seguridad, comprometiendo la integridad, confidencialidad y disponibilidad de la información y los sistemas de la organización.',
            2: 'Las pruebas de vulnerabilidades y evaluaciones de riesgos se realizan de manera limitada o irregular. Se pueden identificar algunas vulnerabilidades, pero otras pueden pasar desapercibidas debido a la falta de exhaustividad y frecuencia en las pruebas. Además, la falta de un enfoque sistemático puede llevar a que las medidas de mitigación no se implementen adecuadamente, aumentando el riesgo de seguridad en la infraestructura de redes.',
            3: 'Las pruebas de vulnerabilidades y evaluaciones de riesgos se realizan regularmente y cubren la mayoría de la infraestructura de redes. Los resultados de estas pruebas se utilizan para mejorar la seguridad, aunque puede haber áreas donde la frecuencia o la profundidad de las pruebas podría ser mejorada. La organización toma medidas correctivas basadas en los resultados, pero puede faltar un proceso formal para asegurar que todas las vulnerabilidades se abordan de manera oportuna y completa.',
            4: 'Las pruebas de vulnerabilidades y evaluaciones de riesgos cumplen totalmente con los requisitos de ISO 27001. Se realizan de manera exhaustiva y regular, cubriendo toda la infraestructura de redes. Los resultados se documentan y se toman medidas correctivas basadas en estos resultados para mitigar cualquier riesgo identificado. Además, se realizan revisiones periódicas de los procedimientos de prueba para asegurar su efectividad y adecuación a las amenazas actuales.',
            5: 'La implementación de pruebas de vulnerabilidades y evaluaciones de riesgos es avanzada y supera los requisitos estándar. Incluye pruebas continuas, técnicas avanzadas de evaluación de riesgos y la utilización de herramientas automatizadas para la detección de vulnerabilidades. La organización está proactivamente mejorando sus métodos de evaluación para mantenerse por delante de las amenazas emergentes. Además, se fomenta una cultura de seguridad que incluye la formación continua del personal en la identificación y mitigación de riesgos.'
        }
    },
    'Control de Acceso a la Información': {
        '¿Se implementan controles para limitar el acceso a la información confidencial y crítica dentro del departamento de sistemas?': {
            1: 'No se implementan controles para limitar el acceso a la información confidencial y crítica, lo que significa que cualquier persona dentro del departamento puede acceder a estos datos sin restricciones. Esto representa un grave riesgo para la seguridad de la información, ya que los datos sensibles pueden ser accedidos, modificados o eliminados sin autorización. La falta de controles de acceso adecuados también implica una ausencia de responsabilidad y supervisión, lo que puede resultar en un uso indebido o malintencionado de la información.',
            2: 'Los controles de acceso a la información confidencial y crítica existen pero se implementan de manera limitada o inconsistente. Algunos datos pueden estar protegidos, pero no todos, y los controles pueden ser fácilmente eludidos. Esto puede resultar en brechas de seguridad y accesos no autorizados. La implementación inconsistente de controles de acceso puede deberse a una falta de políticas claras, a una formación inadecuada del personal o a la ausencia de un monitoreo y revisión regular de los permisos de acceso.',
            3: 'Controles de acceso a la información confidencial y crítica están bien implementados y se utilizan regularmente. Esto incluye la asignación de roles y permisos específicos, y la implementación de políticas de acceso basadas en la necesidad de conocer. Sin embargo, puede haber áreas donde los controles podrían ser mejorados, como en la frecuencia de las revisiones de permisos, la actualización de los accesos en función de cambios en las responsabilidades del personal o la implementación de técnicas de monitoreo más avanzadas para detectar accesos no autorizados.',
            4: 'Los controles de acceso a la información confidencial y crítica cumplen totalmente con los requisitos de ISO 27001. Todos los datos sensibles están protegidos por controles robustos, incluyendo la autenticación multifactor, la segregación de funciones y el monitoreo de accesos. Los controles están documentados y se revisan regularmente para asegurar su efectividad. Además, la organización realiza auditorías periódicas de los permisos de acceso y se asegura de que solo el personal autorizado tenga acceso a la información crítica.',
            5: 'La implementación de controles de acceso a la información confidencial y crítica es avanzada y supera los requisitos estándar. Esto incluye el uso de tecnologías avanzadas como el análisis de comportamiento y la inteligencia artificial para monitorear y limitar los accesos, así como la revisión continua y la mejora de los controles basados en las últimas amenazas y mejores prácticas. La organización también promueve una cultura de seguridad, asegurando que todos los empleados comprendan la importancia de los controles de acceso y se adhieran a las políticas establecidas. Se realiza una formación continua del personal sobre los riesgos asociados a la falta de controles adecuados y las mejores prácticas para proteger la información crítica.'
        },
        '¿Se establecen y mantienen políticas para la clasificación y etiquetado de la información dentro del departamento de sistemas?': {
            1: 'No se establecen ni mantienen políticas para la clasificación y etiquetado de la información, lo que significa que no hay un sistema para identificar y gestionar adecuadamente los diferentes tipos de información. Esto puede llevar a una gestión inadecuada de los datos, comprometiendo su seguridad y disponibilidad. La falta de clasificación y etiquetado también puede resultar en el acceso no autorizado a información confidencial, ya que no hay directrices claras sobre cómo debe ser manejada.',
            2: 'Las políticas de clasificación y etiquetado existen pero no se mantienen adecuadamente. Pueden haber sido establecidas en algún momento, pero no se actualizan para reflejar cambios en la información o en los requisitos de seguridad. Esto puede resultar en inconsistencias y en una clasificación inadecuada de los datos. La falta de actualización y mantenimiento regular de las políticas puede llevar a que la información crítica no esté debidamente protegida, exponiendo a la organización a riesgos innecesarios.',
            3: 'Las políticas de clasificación y etiquetado están bien implementadas y se mantienen regularmente. Esto incluye la identificación y etiquetado de la información según su nivel de sensibilidad y la aplicación de controles de acceso correspondientes. Sin embargo, puede haber áreas donde las políticas podrían ser mejoradas, como en la capacitación del personal sobre la importancia de la clasificación y el etiquetado, o en la integración de nuevas tecnologías que faciliten la clasificación automática de la información.',
            4: 'Las políticas de clasificación y etiquetado cumplen totalmente con los requisitos de ISO 27001. Todos los datos están correctamente clasificados y etiquetados según su nivel de sensibilidad, y se aplican controles adecuados para proteger la información. Las políticas están documentadas y se revisan regularmente para asegurar su efectividad. Además, la organización realiza auditorías periódicas para verificar la correcta clasificación y etiquetado de la información, asegurando que se mantengan los niveles adecuados de protección.',
            5: 'La implementación de políticas de clasificación y etiquetado es avanzada y supera los requisitos estándar. Esto incluye la utilización de tecnologías avanzadas para la clasificación automática de datos, la revisión continua y la mejora de las políticas basadas en las últimas mejores prácticas y amenazas emergentes. La organización está proactivamente mejorando estos procesos para asegurar una gestión óptima de la información. Además, se fomenta una cultura de concienciación sobre la clasificación y etiquetado de la información, proporcionando formación continua y recursos para mantener la adherencia a las políticas establecidas.'
        }
    },
    'Gestión de Incidentes de Seguridad de la Información': {
        '¿Existe un procedimiento documentado para la gestión de incidentes de seguridad de la información?': {
            1: 'No hay un procedimiento documentado para la gestión de incidentes de seguridad, lo que significa que no hay un plan establecido para identificar, responder y mitigar los incidentes de seguridad. Esto puede resultar en una respuesta inadecuada y descoordinada ante incidentes, comprometiendo la seguridad de la información. La falta de un procedimiento documentado también implica que no hay claridad sobre las responsabilidades y acciones a tomar en caso de un incidente, lo que puede aumentar el impacto y la duración del mismo.',
            2: 'El procedimiento para la gestión de incidentes existe pero no está actualizado o se implementa de manera limitada. Puede haber un documento antiguo que no refleja los cambios recientes en la infraestructura de TI o en las amenazas de seguridad. Esto puede llevar a inconsistencias y a una respuesta ineficaz ante incidentes. Además, la falta de actualización regular del procedimiento puede significar que no se están utilizando las mejores prácticas actuales ni las tecnologías más avanzadas para la gestión de incidentes.',
            3: 'El procedimiento para la gestión de incidentes está bien documentado y se revisa regularmente. Incluye procesos para la identificación, respuesta y mitigación de incidentes de seguridad. Sin embargo, puede haber áreas donde el procedimiento podría ser mejorado, como en la integración de nuevas tecnologías, la capacitación del personal sobre su aplicación y la realización de ejercicios de simulación para probar la efectividad del procedimiento en situaciones reales.',
            4: 'El procedimiento para la gestión de incidentes cumple totalmente con los requisitos de ISO 27001. Está documentado, actualizado y se revisa periódicamente para asegurar su efectividad. Incluye todas las fases necesarias para la gestión de incidentes, desde la identificación y notificación hasta la recuperación y el análisis post-incidente. Además, la organización realiza auditorías periódicas del procedimiento para asegurar que se mantenga alineado con las mejores prácticas y que esté preparado para abordar cualquier tipo de incidente de seguridad.',
            5: 'La implementación del procedimiento para la gestión de incidentes es avanzada y supera los requisitos estándar. Incluye el uso de tecnologías avanzadas para la detección y respuesta a incidentes, la automatización de procesos y la integración de inteligencia artificial para mejorar la eficiencia y efectividad. La organización está proactivamente mejorando el procedimiento basado en lecciones aprendidas y en las últimas amenazas y mejores prácticas. También se realizan simulaciones regulares para probar y ajustar el procedimiento, asegurando una respuesta óptima y coordinada ante cualquier incidente de seguridad.'
        },
        '¿Se realiza capacitación y simulacros periódicos para el personal sobre cómo responder a incidentes de seguridad de la información?': {
            1: 'No se realizan capacitaciones ni simulacros sobre incidentes de seguridad, lo que significa que el personal no está preparado para responder adecuadamente a incidentes de seguridad. Esto puede resultar en una respuesta ineficaz y descoordinada, comprometiendo la capacidad de la organización para mitigar los impactos de los incidentes. La falta de formación y simulacros también indica una falta de concienciación sobre la importancia de la preparación y respuesta ante incidentes, lo que puede aumentar el riesgo de daños significativos.',
            2: 'Las capacitaciones y simulacros se realizan de manera irregular o insuficiente. Pueden haber algunas sesiones de capacitación, pero no son frecuentes ni abarcativas. Esto puede llevar a una falta de preparación y a una respuesta inadecuada del personal ante incidentes de seguridad. La formación inconsistente puede significar que solo algunos empleados están bien preparados, mientras que otros no saben cómo actuar en caso de un incidente, lo que compromete la respuesta global de la organización.',
            3: 'Las capacitaciones y simulacros se realizan regularmente, cumpliendo con la mayoría de los requisitos de ISO 27001. El personal está generalmente preparado para responder a incidentes de seguridad, aunque puede haber áreas donde se podrían realizar mejoras adicionales, como en la frecuencia o en la profundidad de los simulacros. Además, es posible que no todos los escenarios de incidentes estén cubiertos en los ejercicios, lo que puede dejar algunas áreas de la respuesta sin probar.',
            4: 'Las capacitaciones y simulacros cumplen totalmente con los requisitos de ISO 27001. Se realizan de manera regular y abarcativa, asegurando que todo el personal esté bien preparado para responder a incidentes de seguridad. Los simulacros se documentan y se revisan para identificar áreas de mejora. Además, la organización asegura que los empleados no solo conozcan los procedimientos, sino que también comprendan la importancia de su rol en la respuesta a incidentes.',
            5: 'La implementación de capacitaciones y simulacros es avanzada y supera los requisitos estándar. Incluye la realización de simulacros detallados y frecuentes, la utilización de escenarios realistas y la integración de nuevas técnicas y tecnologías. La organización está proactivamente mejorando la capacitación del personal y la preparación para incidentes basándose en lecciones aprendidas y en las últimas amenazas y mejores prácticas. Además, se fomenta una cultura de seguridad en toda la organización, asegurando que todos los empleados comprendan la importancia de estar preparados y sepan cómo actuar rápidamente y de manera efectiva ante cualquier incidente de seguridad.'
        }
    }
}

# Procesar las calificaciones y calcular los promedios
def procesar_calificaciones(calificaciones):
    promedios = {aspecto: sum(valores[1] for valores in lista) / len(lista) for aspecto, lista in calificaciones.items()}
    promedios_ponderados = {aspecto: (promedio / 5) * 20 for aspecto, promedio in promedios.items()}

    calificacion_final = sum(promedios_ponderados.values()) / len(promedios_ponderados) * 5
    return promedios_ponderados, calificacion_final

# Generar gráfico de barras utilizando matplotlib
def generar_grafico(promedios_ponderados):
    aspectos = list(promedios_ponderados.keys())
    valores = list(promedios_ponderados.values())

    fig, ax = plt.subplots(figsize=(10, 6), dpi=200)  # Increased DPI for better resolution
    ax.barh(aspectos, valores, color='skyblue')
    ax.set_xlabel('Nivel de Cumplimiento (sobre 20)')
    ax.set_title('Gráfico de Nivel de Cumplimiento por Aspecto')
    ax.set_xlim(0, 20)

    plt.tight_layout()  # Adjust layout to ensure everything fits

    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    return buf

# Generar gráfico de radar utilizando matplotlib
def generar_grafico_radar(promedios_ponderados):
    etiquetas = list(promedios_ponderados.keys())
    valores = list(promedios_ponderados.values())
    valores += valores[:1]  # Añadir el primer valor al final para cerrar el gráfico

    angulos = np.linspace(0, 2 * np.pi, len(etiquetas), endpoint=False).tolist()
    angulos += angulos[:1]

    fig, ax = plt.subplots(figsize=(6, 6), dpi=200, subplot_kw=dict(polar=True))  # Increased DPI
    ax.fill(angulos, valores, color='skyblue', alpha=0.25)
    ax.plot(angulos, valores, color='skyblue', linewidth=2)
    ax.set_yticklabels([])
    ax.set_xticks(angulos[:-1])
    ax.set_xticklabels(etiquetas, fontsize=8)  # Reduced font size for labels
    ax.set_title('Gráfico de Radar por Aspecto', fontsize=10)

    plt.tight_layout()  # Adjust layout to ensure everything fits

    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    return buf

# Generar la conclusión general basada en la calificación final
def generar_conclusion(calificacion_final):
    if 0 <= calificacion_final <= 25:
        return ("El departamento de sistemas muestra una falta significativa de cumplimiento en la gestión de acceso, "
                "seguridad física y ambiental, gestión de comunicaciones y operaciones, control de acceso a la información, "
                "y gestión de incidentes de seguridad de la información. No existen políticas ni procedimientos documentados, "
                "y los controles de seguridad son insuficientes o inexistentes, exponiendo la información a riesgos severos. "
                "Esta deficiencia indica una falta de alineación con los requisitos básicos de la norma ISO 27001, incluyendo "
                "la ausencia de controles formales, falta de gestión de riesgos, y ausencia de procedimientos de respuesta a incidentes. "
                "La falta de políticas y procedimientos documentados significa que no hay una estructura definida para la gestión "
                "de la seguridad de la información, lo que expone a la organización a una alta probabilidad de incidentes de seguridad "
                "graves y a la incapacidad de responder adecuadamente a ellos.")
    elif 26 <= calificacion_final <= 50:
        return ("El departamento de sistemas tiene algunos controles y políticas en su lugar, pero estos no son suficientemente robustos "
                "o no se aplican consistentemente. Existen políticas y procedimientos documentados en algunas áreas, pero pueden estar "
                "desactualizados o no ser efectivos en la práctica. Los controles de seguridad física y ambiental, así como las medidas "
                "de autenticación, se implementan de manera limitada, y las pruebas de vulnerabilidad y evaluaciones de riesgos se realizan "
                "de forma irregular. Esto indica que, aunque hay una conciencia de la necesidad de seguridad de la información, la implementación "
                "no es suficiente para cumplir con los estándares de la norma ISO 27001. La falta de consistencia y actualización en las políticas "
                "y procedimientos sugiere que la organización puede no estar preparada para manejar nuevos riesgos o cambios en el entorno tecnológico, "
                "lo que deja brechas significativas en la seguridad.")
    elif 51 <= calificacion_final <= 75:
        return ("El departamento de sistemas ha implementado la mayoría de los controles de seguridad requeridos por la norma ISO 27001. "
                "Las políticas y procedimientos están documentados y se revisan regularmente. La seguridad física y ambiental es adecuada, "
                "y los controles de autenticación son robustos para la mayoría de los sistemas críticos. Las pruebas de vulnerabilidad y "
                "evaluaciones de riesgos se realizan de manera regular, aunque aún existen áreas que pueden mejorarse para alcanzar un nivel óptimo. "
                "Este nivel de cumplimiento muestra un compromiso con la seguridad de la información y una estructura que, en general, está alineada "
                "con los requisitos de ISO 27001. Sin embargo, para alcanzar un nivel de seguridad óptimo, se deben abordar las áreas que aún necesitan "
                "mejora, como la implementación de controles adicionales y la mejora continua de los procesos existentes.")
    elif 76 <= calificacion_final <= 100:
        return ("El departamento de sistemas cumple completamente con los requisitos de la norma ISO 27001, y además implementa medidas adicionales "
                "que superan los estándares establecidos. Las políticas y procedimientos están completamente documentados y actualizados, y se revisan "
                "periódicamente. La seguridad física y ambiental es robusta y se monitorea continuamente. Los controles de autenticación incluyen medidas "
                "avanzadas como la autenticación multifactor, y las pruebas de vulnerabilidad y evaluaciones de riesgos se realizan de manera continua y "
                "exhaustiva. La gestión de incidentes es proactiva, con simulacros y capacitaciones regulares que aseguran una preparación adecuada del personal. "
                "Este nivel de cumplimiento no solo muestra que la organización cumple con todos los requisitos de ISO 27001, sino que también adopta un enfoque "
                "proactivo para mejorar continuamente su postura de seguridad. La implementación de controles avanzados y la cultura de seguridad establecida "
                "dentro de la organización indican una alta madurez en la gestión de la seguridad de la información.")
    else:
        return "Calificación no válida."

# Generar el informe en Word con las nuevas mejoras
def generar_informe_word(calificaciones, promedios_ponderados, calificacion_final, nombre_auditor, nombre_compania, fecha_evaluacion, nombre_compania_evaluada, destinatario, firma):
    document = Document()

    # Carátula
    document.add_heading('Informe de Evaluación de Cumplimiento de la Norma ISO 27001 (Sistema de Gestión de Seguridad de la Información)', 0)
    document.add_paragraph(f'Compañía Evaluada: {nombre_compania_evaluada}', style='Title')
    document.add_paragraph(f'Compañía Auditora: {nombre_compania}', style='Heading 3')
    document.add_paragraph(f'Auditor: {nombre_auditor}', style='Heading 3')
    document.add_paragraph(f'Fecha de Evaluación: {fecha_evaluacion}', style='Heading 3')
    
    # Añadir un salto de página
    document.add_page_break()

    # Índice
    document.add_heading('Índice', level=1)
    document.add_paragraph("1. Carta de Introducción")
    document.add_paragraph("2. Limitación de Responsabilidad")
    document.add_paragraph("3. Objetivo de la Norma ISO 27001")
    document.add_paragraph("4. Dimensiones Evaluadas")
    document.add_paragraph("5. Metodología de Calificación")
    document.add_paragraph("6. Resultados de la Evaluación")
    document.add_paragraph("7. Conclusión General")
    document.add_paragraph("8. Gráfico de Nivel de Cumplimiento por Aspecto")
    document.add_paragraph("9. Gráfico de Radar por Aspecto")

    # Añadir un salto de página
    document.add_page_break()

    # Carta de introducción
    document.add_heading('Carta de Introducción', level=1)
    document.add_paragraph(
        f"Estimado/a {destinatario},\n\n"
        "A través de la presente, tenemos el agrado de presentar los resultados de la evaluación realizada en su organización en relación con el cumplimiento del Sistema de Gestión de Seguridad de la Información (SGSI) conforme a la norma ISO 27001. La evaluación fue llevada a cabo con el objetivo de revisar y analizar la efectividad de las políticas, procedimientos y controles implementados para asegurar que se alineen con los estándares internacionales de seguridad de la información.\n\n"
        "El informe adjunto contiene un resumen detallado de las áreas revisadas, incluyendo gestión de acceso, seguridad física y ambiental, gestión de comunicaciones y operaciones, control de acceso a la información, y gestión de incidentes de seguridad de la información. Además, se ofrecen recomendaciones para la optimización de los procesos y la mejora continua del SGSI.\n\n"
        "Esperamos que los resultados presentados en este informe sean de utilidad para fortalecer las prácticas de seguridad de la información en su organización. Quedamos a su disposición para profundizar en cualquier aspecto del informe que requiera su atención.\n\n"
        f"Atentamente,\n\n{firma}"
    )

    # Añadir Limitación de Responsabilidad
    document.add_page_break()
    document.add_heading('Limitación de Responsabilidad', level=1)
    document.add_paragraph(
        "La presente evaluación ha sido realizada sobre la base de la información proporcionada por la organización evaluada y las observaciones efectuadas durante el proceso de evaluación. Si bien se ha aplicado la debida diligencia y se han seguido los estándares reconocidos para llevar a cabo esta revisión, los resultados y recomendaciones contenidos en este informe no garantizan la seguridad absoluta del sistema de gestión evaluado. La responsabilidad de implementar, mantener y mejorar los controles de seguridad recae exclusivamente en la organización evaluada. La compañía evaluadora no asume responsabilidad alguna por los resultados derivados de la implementación o falta de implementación de las recomendaciones sugeridas en este informe."
    )

    # Añadir un salto de página
    document.add_page_break()

    # Descripción del objetivo de la norma
    document.add_heading('Objetivo de la Norma ISO 27001', level=1)
    document.add_paragraph(
        "La norma ISO/IEC 27001 establece los requisitos para un sistema de gestión de seguridad de la información (SGSI), "
        "incluyendo los aspectos relacionados con la implementación, el mantenimiento y la mejora continua del SGSI. "
        "Su objetivo es proteger la información dentro de la organización, asegurando su confidencialidad, integridad y disponibilidad."
    )

    # Descripción de las dimensiones evaluadas
    document.add_heading('Dimensiones Evaluadas', level=1)
    document.add_paragraph(
        "A continuación se detallan las diferentes dimensiones evaluadas en este informe, junto con una breve descripción de cada una:"
    )

    dimensiones = {
        'Gestión de Acceso': "Evalúa la existencia y eficacia de políticas y procedimientos para la gestión de accesos, "
                             "incluyendo controles de autenticación y autorización para proteger los sistemas críticos.",
        'Seguridad Física y Ambiental': "Evalúa las medidas de seguridad física y controles ambientales implementados para proteger "
                                        "los equipos e infraestructuras críticas de la organización.",
        'Gestión de Comunicaciones y Operaciones': "Evalúa los procedimientos seguros para la transmisión de datos sensibles y las prácticas "
                                                   "de gestión de operaciones para mantener la seguridad de la infraestructura de red.",
        'Control de Acceso a la Información': "Evalúa los controles implementados para limitar el acceso a la información confidencial y crítica, "
                                              "así como las políticas de clasificación y etiquetado de la información.",
        'Gestión de Incidentes de Seguridad de la Información': "Evalúa la existencia y eficacia de procedimientos para la gestión de incidentes "
                                                               "de seguridad, incluyendo la capacitación y los simulacros realizados para preparar al personal."
    }

    for dimension, descripcion in dimensiones.items():
        document.add_heading(dimension, level=2)
        document.add_paragraph(descripcion)

    # Metodología de calificación
    document.add_heading('Metodología de Calificación', level=1)
    document.add_paragraph(
        "La evaluación se basa en una escala de 1 a 5, donde cada valor representa el nivel de cumplimiento de la norma:"
    )
    calificacion_metodologia = {
        1: "1 = No Cumple: No se realiza ninguna acción o la acción es insuficiente.",
        2: "2 = Cumple Parcialmente: Las acciones se realizan pero no con la frecuencia o efectividad requerida.",
        3: "3 = Cumple en Gran Medida: Las acciones se realizan regularmente y cumplen con la mayoría de los requisitos.",
        4: "4 = Cumple Totalmente: Las acciones cumplen con todos los requisitos establecidos.",
        5: "5 = Cumple y Supera las Expectativas: Se implementan medidas adicionales que superan los requisitos establecidos."
    }

    for key, value in calificacion_metodologia.items():
        document.add_paragraph(value)

    # Resultados de la evaluación
    document.add_heading('Resultados de la Evaluación', level=1)
    for aspecto, preguntas in calificaciones.items():
        document.add_heading(aspecto, level=2)
        for pregunta, calificacion in preguntas:
            descripcion = rubricas[aspecto][pregunta][calificacion]
            p = document.add_paragraph()
            p.add_run(f'{pregunta}: ').bold = True
            p.add_run(f'{calificacion} - {descripcion}')
        document.add_paragraph(f'Promedio del aspecto ({aspecto}): {promedios_ponderados[aspecto]:.2f} / 20')
        document.add_paragraph()

    document.add_paragraph(f'Calificación final del departamento de sistemas: {calificacion_final:.2f} / 100')
    document.add_paragraph()

    # Conclusión general
    conclusion = generar_conclusion(calificacion_final)
    document.add_heading('Conclusión General', level=1)
    document.add_paragraph(conclusion)
    document.add_paragraph()

    # Añadir gráficos generados en memoria
    document.add_heading('Gráfico de Nivel de Cumplimiento por Aspecto', level=1)
    buf_barras = generar_grafico(promedios_ponderados)
    document.add_picture(buf_barras, width=Inches(6))

    document.add_heading('Gráfico de Radar por Aspecto', level=1)
    buf_radar = generar_grafico_radar(promedios_ponderados)
    document.add_picture(buf_radar, width=Inches(6))

    # Añadir pie de página
    section = document.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f'Compañía Auditora: {nombre_compania} - Fecha de Evaluación: {fecha_evaluacion}'
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document

# Interfaz en Streamlit
st.title("Evaluación de Cumplimiento ISO 27001")

# Descripción ISO 27001
st.write("""
ISO 27001 es una norma internacional que establece los requisitos fundamentales para un sistema de gestión de seguridad de la información (SGSI). Su propósito es garantizar la protección de la información sensible de una organización, asegurando su confidencialidad, integridad y disponibilidad mediante la implementación de políticas y controles bien definidos.

A través de esta aplicación, tendrás la capacidad de evaluar de manera estructurada diversos aspectos críticos de la seguridad de la información en la organización que estás auditando. La interfaz intuitiva te guiará a través de cada sección, permitiéndote ingresar la información relevante y calificar las prácticas actuales de la organización de manera clara y precisa.
""")

# Datos generales
st.header("Parte 1: Datos Generales")
st.write("""
Para comenzar, por favor completa la información clave que personalizará y contextualizará el informe de auditoría. Asegúrate de ingresar los detalles necesarios tanto de la entidad auditora como de la organización evaluada. Además, incluye la fecha en que se realiza la evaluación y los datos de contacto de la persona que recibirá el informe. Con estos datos, garantizamos que el informe refleje fielmente el proceso de auditoría realizado.
""")
nombre_auditor = st.text_input("Nombre del Auditor")
nombre_compania = st.text_input("Nombre de la Compañía Auditora")
nombre_compania_evaluada = st.text_input("Nombre de la Compañía Evaluada")
fecha_evaluacion = st.date_input("Fecha de Evaluación")
destinatario = st.text_input("Destinatario del Informe")
firma = st.text_input("Firma del Auditor")

# Evaluación por aspectos
st.header("Parte 2: Evaluación por Aspectos")
calificaciones_input = {key: [] for key in rubricas.keys()}

# Instrucción sobre la evaluación
st.write("""
Para cada uno de los cinco aspectos clave de la seguridad de la información, asigna una calificación en una escala del 1 al 5, donde 1 representa el nivel más bajo de cumplimiento y 5 el nivel más alto. Esta evaluación es esencial para asegurar que el diagnóstico sea preciso y útil para la organización en su proceso de alineación con los estándares de la norma ISO 27001.
""")

for aspecto, preguntas in rubricas.items():
    st.subheader(aspecto)
    for pregunta, opciones in preguntas.items():
        calificacion = st.selectbox(pregunta, list(opciones.keys()), format_func=lambda x: f"{x}: {opciones[x]}")
        calificaciones_input[aspecto].append((pregunta, calificacion))

# Párrafo para generación del informe
st.header("Parte 3: Generar el informe")
st.write("""
Al concluir la evaluación, la herramienta generará automáticamente un informe detallado que resumirá los resultados de la auditoría. Este informe te proporcionará un diagnóstico completo del estado actual de la seguridad en la organización, junto con recomendaciones específicas para mejorar el cumplimiento de los estándares de ISO 27001. Además, tendrás acceso a gráficos interactivos que te ayudarán a interpretar los datos de manera visual y a identificar fácilmente las áreas que requieren mayor atención.
""")

# Botón para generar el informe
if st.button("Generar Informe"):
    if not all([nombre_auditor, nombre_compania, nombre_compania_evaluada, fecha_evaluacion, destinatario, firma]):
        st.error("Debe completar todos los campos para generar el informe.")
    else:
        promedios_ponderados, calificacion_final = procesar_calificaciones(calificaciones_input)
        document = generar_informe_word(calificaciones_input, promedios_ponderados, calificacion_final,
                                        nombre_auditor, nombre_compania, fecha_evaluacion,
                                        nombre_compania_evaluada, destinatario, firma)
        # Guardar el archivo en un buffer de memoria
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        st.download_button(
            label="Descargar Informe en Word",
            data=buf,
            file_name="informe_ISO27001.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
